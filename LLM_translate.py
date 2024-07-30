import docx
from docx import Document
import os
import logging
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import configparser
from pathlib import Path
import requests
from requests.exceptions import Timeout
import json
import queue
import threading
from threading import Thread
import time 
import sys
import re
from copy import deepcopy
from urllib.parse import urlparse
from requests.exceptions import Timeout, RequestException
from tenacity import retry, stop_after_attempt, wait_exponential


logging.basicConfig(filename='translation_log.txt', level=logging.DEBUG)
CONFIG_FILE = 'translation_config.ini'

LLM_response_timeout_s = 20

def load_config():
    config = configparser.ConfigParser()
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
    return config

def save_config(config):
    with open(CONFIG_FILE, 'w') as configfile:
        config.write(configfile)

def contains_words(text):
    return len(text.strip()) > 1 and bool(re.search(r'\w', text))

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def make_api_request(api_url, headers, payload, timeout):
    response = requests.post(api_url, headers=headers, json=payload, stream=True, timeout=timeout)
    response.raise_for_status()
    return response

def translate_text(text, api_url, api_key, target_language, model, timeout=LLM_response_timeout_s, shutdown_flag=None):
    if not contains_words(text):
        return text  # Return the original text if it doesn't contain any word characters
    if shutdown_flag and shutdown_flag.is_set():
        return None

    # Determine API type based on URL
    parsed_url = urlparse(api_url)
    is_ollama = parsed_url.hostname in ['localhost', '127.0.0.1']

    prompt = f"INPUT TEXT:{text} END OF INPUT TEXT. Translate the previous text to {target_language} language. Only return the translation, no other text. No explanation."

    headers = {
        "Content-Type": "application/json"
    }

    if not is_ollama:
        headers["Authorization"] = f"Bearer {api_key}"

    if is_ollama:
        payload = {
            "model": model,
            "prompt": prompt
        }
    else:  # OpenRouter
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}]
        }

    try:
        response = make_api_request(api_url, headers, payload, timeout)
        
        # Check if the response is a single JSON object
        content_type = response.headers.get('Content-Type', '')
        if 'application/json' in content_type:
            json_response = response.json()
            if is_ollama:
                full_response = json_response.get('response', '')
            else:  # OpenRouter
                full_response = json_response.get('choices', [{}])[0].get('message', {}).get('content', '')
        else:
            # Handle streaming response
            full_response = ""
            start_time = time.time()
            for line in response.iter_lines():
                if time.time() - start_time > timeout:
                    raise requests.exceptions.Timeout("Response streaming timed out")
                if line:
                    json_response = json.loads(line)
                    if is_ollama:
                        if 'response' in json_response:
                            full_response += json_response['response']
                        if json_response.get('done', False):
                            break
                    else:  # OpenRouter
                        if 'choices' in json_response and json_response['choices']:
                            full_response += json_response['choices'][0]['message']['content']
                        if json_response.get('done', False):
                            break
                # Reset the timer after each successful line
                start_time = time.time()
        
        print(full_response.strip())
        return full_response.strip()

    except Timeout:
        error_message = "API request timed out"
        logging.error(error_message)
        raise
    except requests.exceptions.RequestException as e:
        error_message = f"Request error: {str(e)}"
        logging.error(error_message)
        raise
    except json.JSONDecodeError as e:
        error_message = f"JSON decode error: {str(e)}\nResponse content: {response.text}"
        logging.error(error_message)
        raise
    except Exception as e:
        error_message = f"An unexpected error occurred: {str(e)}"
        logging.error(error_message)
        raise

def preserve_run_formatting(new_run, original_run):
    # Preserve character formatting
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic
    new_run.underline = original_run.underline
    new_run.font.strike = original_run.font.strike
    new_run.font.subscript = original_run.font.subscript
    new_run.font.superscript = original_run.font.superscript
    new_run.font.size = original_run.font.size
    new_run.font.color.rgb = original_run.font.color.rgb
    new_run.font.highlight_color = original_run.font.highlight_color
    new_run.font.name = original_run.font.name
    new_run.style = original_run.style

def preserve_paragraph_formatting(new_paragraph, original_paragraph):
    # Preserve paragraph formatting
    new_paragraph.alignment = original_paragraph.alignment
    new_paragraph.style = original_paragraph.style
    new_paragraph.paragraph_format.left_indent = original_paragraph.paragraph_format.left_indent
    new_paragraph.paragraph_format.right_indent = original_paragraph.paragraph_format.right_indent
    new_paragraph.paragraph_format.first_line_indent = original_paragraph.paragraph_format.first_line_indent
    new_paragraph.paragraph_format.line_spacing = original_paragraph.paragraph_format.line_spacing
    new_paragraph.paragraph_format.space_before = original_paragraph.paragraph_format.space_before
    new_paragraph.paragraph_format.space_after = original_paragraph.paragraph_format.space_after
    new_paragraph.paragraph_format.keep_together = original_paragraph.paragraph_format.keep_together
    new_paragraph.paragraph_format.keep_with_next = original_paragraph.paragraph_format.keep_with_next
    new_paragraph.paragraph_format.page_break_before = original_paragraph.paragraph_format.page_break_before

def translate_paragraph(paragraph, api_url, api_key, target_language, model):
    if paragraph.text.strip():
        translated_text = translate_text(paragraph.text, api_url, api_key, target_language, model)
        new_paragraph = deepcopy(paragraph)
        new_paragraph.clear()
        preserve_paragraph_formatting(new_paragraph, paragraph)
        
        new_run = new_paragraph.add_run(translated_text)
        preserve_run_formatting(new_run, paragraph.runs[0] if paragraph.runs else None)
        
        paragraph._p.getparent().replace(paragraph._p, new_paragraph._p)
        return True
    return False

def translate_table(table, api_url, api_key, target_language, model, progress_callback, processed_elements, total_elements, shutdown_flag):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if shutdown_flag.is_set():
                    return processed_elements
                if translate_paragraph(paragraph, api_url, api_key, target_language, model):
                    processed_elements += 1
                    progress_callback(int(processed_elements / total_elements * 100), f"Translating table text: {paragraph.text[:300]}...")
    return processed_elements

def translate_document(input_file, output_file, api_url, api_key, target_language, model, progress_callback, shutdown_flag):
    input_file = Path(input_file)
    output_file = Path(output_file)

    try:
        doc = Document(input_file)

        # Count total translatable elements
        total_elements = sum(1 for paragraph in doc.paragraphs if paragraph.text.strip())
        total_elements += sum(1 for table in doc.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs if paragraph.text.strip())
        total_elements += sum(1 for shape in doc.inline_shapes if shape.has_text_frame for paragraph in shape.text_frame.paragraphs if paragraph.text.strip())
        for section in doc.sections:
            total_elements += sum(1 for paragraph in section.header.paragraphs if paragraph.text.strip())
            total_elements += sum(1 for paragraph in section.footer.paragraphs if paragraph.text.strip())

        processed_elements = 0

        if not shutdown_flag.is_set():
            # Translate main document text
            for paragraph in doc.paragraphs:
                if shutdown_flag.is_set():
                    return
                if translate_paragraph(paragraph, api_url, api_key, target_language, model):
                    processed_elements += 1
                    progress_callback(int(processed_elements / total_elements * 100), f"Translating: {paragraph.text[:300]}...")

            # Translate text in tables
            for table in doc.tables:
                processed_elements = translate_table(table, api_url, api_key, target_language, model, progress_callback, processed_elements, total_elements, shutdown_flag)

            # Translate text in text boxes (shapes)
            for shape in doc.inline_shapes:
                if shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        if shutdown_flag.is_set():
                            return
                        if translate_paragraph(paragraph, api_url, api_key, target_language, model):
                            processed_elements += 1
                            progress_callback(int(processed_elements / total_elements * 100), f"Translating shape text: {paragraph.text[:300]}...")

            # Translate headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    if shutdown_flag.is_set():
                        return
                    if translate_paragraph(paragraph, api_url, api_key, target_language, model):
                        processed_elements += 1
                        progress_callback(int(processed_elements / total_elements * 100), f"Translating header: {paragraph.text[:300]}...")
                
                for paragraph in section.footer.paragraphs:
                    if shutdown_flag.is_set():
                        return
                    if translate_paragraph(paragraph, api_url, api_key, target_language, model):
                        processed_elements += 1
                        progress_callback(int(processed_elements / total_elements * 100), f"Translating footer: {paragraph.text[:300]}...")

            doc.save(output_file)
            progress_callback(100, "Translation completed!")
            print(f"Successfully translated and saved: {output_file}")
        else:  # if shutdown detected
            return

    except Exception as e:
        error_message = f"Error in translate_document: {str(e)}"
        logging.error(error_message)
        raise

class TranslationGUI:
    def __init__(self, master):
        self.master = master
        self.master.title("Document Translator")
        self.master.geometry("1000x600") 
        self.config = load_config()

        # To handle shutdown of running threads upon GUI close
        self.shutdown_flag = threading.Event()
        self.master.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.create_widgets()
        self.load_api_list()
        self.load_target_language()
        self.api_tree.bind('<<TreeviewSelect>>', self.on_api_select)

        # Add progress bar and status label
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(self.master, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(padx=10, pady=5, fill=tk.X)
        
        self.status_var = tk.StringVar()
        self.status_label = ttk.Label(self.master, textvariable=self.status_var, wraplength = 980)
        self.status_label.pack(padx=10, pady=5)

        self.queue = queue.Queue()
        self.check_queue()
    
    def test_api(self):
        selection = self.api_tree.selection()
        if not selection:
            messagebox.showerror("Error", "Please select an API to test")
            return

        item = selection[0]
        name, model, url, key = self.api_tree.item(item, 'values')

        test_prompt = "Translate the following English text to French: 'Hello, world!'"
        
        try:
            self.status_var.set("Testing API connection...")
            self.master.update_idletasks()

            result = translate_text(test_prompt, url, key, "French", model, timeout=15)  # 15 seconds timeout for the test
            
            if result:
                messagebox.showinfo("API Test Successful", f"API connection successful!\nResponse: {result}")
            else:
                messagebox.showerror("API Test Failed", "API connection failed. No response received.")
        except Timeout:
            messagebox.showerror("API Test Failed", "API request timed out. The server took too long to respond.")
        except Exception as e:
            error_message = f"API Test Failed: {str(e)}"
            logging.error(error_message)
            messagebox.showerror("API Test Failed", error_message)
        finally:
            self.status_var.set("")
            self.master.update_idletasks()

    def on_closing(self):
        self.shutdown_flag.set()
        self.master.after(3000, self.master.quit)

    def update_progress(self, value, status):
        self.progress_var.set(value)
        self.status_var.set(status)

    def create_widgets(self):
        # API List
        self.api_frame = ttk.LabelFrame(self.master, text="API Settings")
        self.api_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

        # Replace Listbox with Treeview
        self.api_tree = ttk.Treeview(self.api_frame, columns=('Name', 'Model', 'URL', 'Key'), show='headings')
        self.api_tree.heading('Name', text='Name')
        self.api_tree.heading('Model', text='Model')
        self.api_tree.heading('URL', text='URL')
        self.api_tree.heading('Key', text='Key')
        self.api_tree.column('Name', width=100)
        self.api_tree.column('Model', width=100)
        self.api_tree.column('URL', width=300)
        self.api_tree.column('Key', width=100)
        self.api_tree.pack(side=tk.LEFT, padx=5, pady=5, fill=tk.BOTH, expand=True)

        self.scroll = ttk.Scrollbar(self.api_frame, orient=tk.VERTICAL, command=self.api_tree.yview)
        self.scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.api_tree.configure(yscrollcommand=self.scroll.set)

        self.delete_button = ttk.Button(self.api_frame, text="Delete Selected", command=self.delete_api)
        self.delete_button.pack(side=tk.BOTTOM, padx=5, pady=5)

        # Add API
        self.add_frame = ttk.Frame(self.master)
        self.add_frame.pack(padx=10, pady=5, fill=tk.X)

        ttk.Label(self.add_frame, text="Name:").pack(side=tk.LEFT, padx=5)
        self.name_entry = ttk.Entry(self.add_frame, width=15)
        self.name_entry.pack(side=tk.LEFT, padx=5)

        ttk.Label(self.add_frame, text="Model:").pack(side=tk.LEFT, padx=5)
        self.model_entry = ttk.Entry(self.add_frame, width=15)
        self.model_entry.pack(side=tk.LEFT, padx=5)

        ttk.Label(self.add_frame, text="URL:").pack(side=tk.LEFT, padx=5)
        self.url_entry = ttk.Entry(self.add_frame, width=30)
        self.url_entry.pack(side=tk.LEFT, padx=5)

        ttk.Label(self.add_frame, text="API Key (optional):").pack(side=tk.LEFT, padx=5)
        self.key_entry = ttk.Entry(self.add_frame, width=15, show="*")
        self.key_entry.pack(side=tk.LEFT, padx=5)

        ttk.Button(self.add_frame, text="Add API", command=self.add_api).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.add_frame, text="Edit API", command=self.edit_api).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.add_frame, text="Test API", command=self.test_api).pack(side=tk.LEFT, padx=5)

        # File Selection
        self.file_frame = ttk.Frame(self.master)
        self.file_frame.pack(padx=10, pady=5, fill=tk.X)

        ttk.Label(self.file_frame, text="Input File:").grid(row=0, column=0, padx=5, pady=2, sticky="w")
        self.input_entry = ttk.Entry(self.file_frame, width=105)
        self.input_entry.grid(row=0, column=1, padx=5, pady=2)
        ttk.Button(self.file_frame, text="Browse", command=self.browse_input).grid(row=0, column=2, padx=5, pady=2)

        ttk.Label(self.file_frame, text="Output File:").grid(row=1, column=0, padx=5, pady=2, sticky="w")
        self.output_entry = ttk.Entry(self.file_frame, width=105)
        self.output_entry.grid(row=1, column=1, padx=5, pady=2)

        # Target Language
        ttk.Label(self.file_frame, text="Target Language:").grid(row=2, column=0, padx=5, pady=2, sticky="w")
        self.lang_entry = ttk.Entry(self.file_frame, width=20)
        self.lang_entry.grid(row=2, column=1, padx=5, pady=2, sticky="w")

        # Translation
        self.trans_frame = ttk.Frame(self.master)
        self.trans_frame.pack(padx=10, pady=5)

        self.open_var = tk.BooleanVar()
        ttk.Checkbutton(self.trans_frame, text="Open file when complete", variable=self.open_var).pack(side=tk.LEFT)
        ttk.Button(self.trans_frame, text="Do Translation", command=self.do_translation).pack(side=tk.LEFT, padx=5)

    def load_api_list(self):
        if 'APIs' in self.config:
            for name, details in self.config['APIs'].items():
                model, url, key = details.split(',')
                self.api_tree.insert('', 'end', values=(name, model, url, key))

    def load_target_language(self):
        if 'Settings' in self.config and 'target_language' in self.config['Settings']:
            self.lang_entry.insert(0, self.config['Settings']['target_language'])

    def add_api(self):
        name = self.name_entry.get()
        model = self.model_entry.get()
        url = self.url_entry.get()
        key = self.key_entry.get()
        if name and model and url:  # Key is not checked here
            if 'APIs' not in self.config:
                self.config['APIs'] = {}
            self.config['APIs'][name] = f"{model},{url},{key}"
            save_config(self.config)
            self.api_tree.insert('', 'end', values=(name, model, url, key))
            self.name_entry.delete(0, tk.END)
            self.model_entry.delete(0, tk.END)
            self.url_entry.delete(0, tk.END)
            self.key_entry.delete(0, tk.END)
        else:
            messagebox.showerror("Error", "Please fill Name, Model, and URL fields")

    def edit_api(self):
        selection = self.api_tree.selection()
        if selection:
            item = selection[0]
            old_name = self.api_tree.item(item, 'values')[0]
            new_name = self.name_entry.get()
            model = self.model_entry.get()
            url = self.url_entry.get()
            key = self.key_entry.get()
            if new_name and model and url:  # Key is not checked here
                # Update config
                del self.config['APIs'][old_name]
                self.config['APIs'][new_name] = f"{model},{url},{key}"
                save_config(self.config)
                # Update treeview
                self.api_tree.item(item, values=(new_name, model, url, key))
                messagebox.showinfo("Success", "API updated successfully")
            else:
                messagebox.showerror("Error", "Please fill Name, Model, and URL fields")
        else:
            messagebox.showerror("Error", "Please select an API to edit")

    def delete_api(self):
        selection = self.api_tree.selection()
        if selection:
            item = selection[0]
            name = self.api_tree.item(item, 'values')[0]
            del self.config['APIs'][name]
            save_config(self.config)
            self.api_tree.delete(item)
        else:
            messagebox.showerror("Error", "Please select an API to delete")

    def browse_input(self):
        filename = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
        if filename:
            input_path = Path(filename)
            if input_path.exists():
                self.input_entry.delete(0, tk.END)
                self.input_entry.insert(0, str(input_path))
                # Set default output filename
                output_path = input_path.with_stem(input_path.stem + "OUT")
                self.output_entry.delete(0, tk.END)
                self.output_entry.insert(0, str(output_path))
            else:
                messagebox.showerror("Error", f"Selected file does not exist: {input_path}")

    def show_success(self, message):
        messagebox.showinfo("Success", message)
        if self.open_var.get():
            os.startfile(str(output_file))

    def show_error(self, message):
        messagebox.showerror("Error", message)

    def reset_progress(self):
        self.progress_var.set(0)
        self.status_var.set("")

    def check_queue(self):
        try:
            while True:
                task = self.queue.get_nowait()
                if task[0] == "progress":
                    self.update_progress(task[1], task[2])
                elif task[0] == "success":
                    self.show_success(task[1])
                elif task[0] == "error":
                    self.show_error(task[1])
        except queue.Empty:
            pass
        finally:
            if not self.shutdown_flag.is_set():
                self.master.after(100, self.check_queue)

    def do_translation(self):
        input_file = Path(self.input_entry.get())
        output_file = Path(self.output_entry.get())
        target_language = self.lang_entry.get()
        selection = self.api_tree.selection()
        
        if not input_file or not output_file or not target_language:
            messagebox.showerror("Error", "Please fill all fields")
            return
        
        if not selection:
            messagebox.showerror("Error", "Please select an API")
            return

        if not input_file.exists():
            error_message = f"Input file does not exist: {input_file}"
            logging.error(error_message)
            messagebox.showerror("Error", error_message)
            return

        if not input_file.is_file():
            error_message = f"Input path is not a file: {input_file}"
            logging.error(error_message)
            messagebox.showerror("Error", error_message)
            return

        item = selection[0]
        name, model, url, key = self.api_tree.item(item, 'values')

        # Save target language to config
        if 'Settings' not in self.config:
            self.config['Settings'] = {}
        self.config['Settings']['target_language'] = target_language
        save_config(self.config)

        def queue_callback(value, status):
            self.queue.put(("progress", value, status))

        def translation_thread():
            try:
                translate_document(input_file, output_file, url, key, target_language, model, queue_callback, self.shutdown_flag)
                if not self.shutdown_flag.is_set():
                    self.queue.put(("success", "Translation complete!"))
            except Exception as e:
                if not self.shutdown_flag.is_set():
                    error_message = f"An error occurred: {str(e)}"
                    logging.error(error_message)
                    self.queue.put(("error", error_message))

        thread = Thread(target=translation_thread)
        thread.start()
        self.check_queue()  # Start checking the queue

    def on_api_select(self, event):
        selection = self.api_tree.selection()
        if selection:
            item = selection[0]
            name, model, url, key = self.api_tree.item(item, 'values')
            self.name_entry.delete(0, tk.END)
            self.name_entry.insert(0, name)
            self.model_entry.delete(0, tk.END)
            self.model_entry.insert(0, model)
            self.url_entry.delete(0, tk.END)
            self.url_entry.insert(0, url)
            self.key_entry.delete(0, tk.END)
            self.key_entry.insert(0, key)

if __name__ == "__main__":
    root = tk.Tk()
    app = TranslationGUI(root)
    root.mainloop()
    sys.exit(0)